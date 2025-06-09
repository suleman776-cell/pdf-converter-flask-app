import os
import sqlite3
import uuid
import zipfile
from flask import (Flask, render_template, redirect, url_for, flash, request,
                   session, send_from_directory, after_this_request)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from itsdangerous import URLSafeTimedSerializer
from flask_mail import Mail, Message
from email_validator import validate_email, EmailNotValidError
import fitz  # PyMuPDF
from docx import Document
import nbformat
from dotenv import load_dotenv

# Load environment variables from a .env file
load_dotenv()

app = Flask(__name__)
# Load secret key from environment variable for security
app.secret_key = os.environ.get('SECRET_KEY', 'a-fallback-secret-key-for-development')

# --- Admin credentials loaded from environment variables ---
ADMIN_USERNAME = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_PASSWORD_HASH = generate_password_hash(os.environ.get('ADMIN_PASSWORD', 'adminpass'))

# --- Database setup ---
DB_NAME = 'users.db'

def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                email TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            );
        ''')
        conn.commit()

# --- Flask-Mail Config ---
app.config.update(
    MAIL_SERVER='smtp.gmail.com',
    MAIL_PORT=587,
    MAIL_USE_TLS=True,
    MAIL_USERNAME=os.environ.get('MAIL_USERNAME'),
    MAIL_PASSWORD=os.environ.get('MAIL_APP_PASSWORD')
)
mail = Mail(app)
serializer = URLSafeTimedSerializer(app.secret_key)

# --- User helper functions ---
def get_user_by_email(email):
    with sqlite3.connect(DB_NAME) as conn:
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE email = ?", (email,))
        return c.fetchone()

def get_user_by_username(username):
    with sqlite3.connect(DB_NAME) as conn:
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM users WHERE username = ?", (username,))
        return c.fetchone()

def create_user(username, email, password_hash):
    with sqlite3.connect(DB_NAME) as conn:
        c = conn.cursor()
        c.execute("INSERT INTO users (username, email, password) VALUES (?, ?, ?)",
                  (username, email, password_hash))
        conn.commit()

def update_password(email, new_password_hash):
    with sqlite3.connect(DB_NAME) as conn:
        c = conn.cursor()
        c.execute("UPDATE users SET password = ? WHERE email = ?", (new_password_hash, email))
        conn.commit()

# --- File paths ---
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
TEMP_FOLDER = 'converter_temp'
for folder in (UPLOAD_FOLDER, CONVERTED_FOLDER, TEMP_FOLDER):
    os.makedirs(folder, exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- Routes ---
@app.route('/')
def home():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    return render_template('home.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username'].strip()
        email = request.form['email'].strip()
        password = request.form['password']
        confirm = request.form['confirm_password']

        if not all([username, email, password, confirm]):
            flash('All fields are required.', 'danger')
            return redirect(url_for('signup'))

        try:
            valid = validate_email(email)
            email = valid.email
        except EmailNotValidError as e:
            flash(str(e), 'danger')
            return redirect(url_for('signup'))

        if password != confirm:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('signup'))

        if get_user_by_username(username):
            flash('Username already exists.', 'warning')
            return redirect(url_for('signup'))
        if get_user_by_email(email):
            flash('Email already registered.', 'warning')
            return redirect(url_for('signup'))

        hashed_password = generate_password_hash(password)
        create_user(username, email, hashed_password)
        flash('Registration successful. Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('signup.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        identifier = request.form['email_or_username'].strip()
        password = request.form['password']
        user = get_user_by_email(identifier) or get_user_by_username(identifier)
        
        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['id']
            session['username'] = user['username']
            flash(f"Welcome, {user['username']}!", 'success')
            return redirect(url_for('dashboard'))

        flash('Invalid credentials. Please try again.', 'danger')
        return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))

@app.route('/send_otp', methods=['GET', 'POST'])
def send_otp():
    if request.method == 'POST':
        email = request.form['email'].strip()

        try:
            validate_email(email)
        except EmailNotValidError as e:
            flash(str(e), 'danger')
            return redirect(url_for('send_otp'))
        
        token = serializer.dumps(email, salt='otp-salt')
        otp_code = token[:6].upper()

        session['otp_email'] = email
        session['otp_code'] = otp_code
        
        try:
            msg = Message('Your OTP Code - FlaskConverter',
                          sender=app.config['MAIL_USERNAME'],
                          recipients=[email])
            msg.body = f'Your one-time password (OTP) is: {otp_code}\n\nThis code will expire in 5 minutes.'
            mail.send(msg)
            flash('An OTP has been sent to your email address.', 'info')
            return redirect(url_for('verify_otp'))
        except Exception as e:
            app.logger.error(f"Mail sending failed: {e}")
            flash('Sorry, there was an error sending the OTP. Please try again later.', 'danger')

    return render_template('send_otp.html')


@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():
    if 'otp_email' not in session:
        flash('Please request an OTP first.', 'warning')
        return redirect(url_for('send_otp'))
        
    if request.method == 'POST':
        user_otp = request.form['otp'].strip()
        
        if 'otp_code' in session and user_otp == session.get('otp_code'):
            session.pop('otp_code', None)
            flash('OTP verified successfully. You can now register.', 'success')
            return redirect(url_for('signup'))
        else:
            flash('Invalid or expired OTP. Please try again.', 'danger')
            
    return render_template('verify_otp.html')


@app.route('/reset_password', methods=['GET', 'POST'])
def reset_request():
    if request.method == 'POST':
        email = request.form['email'].strip()
        user = get_user_by_email(email)
        if user:
            token = serializer.dumps(email, salt='reset-salt')
            link = url_for('reset_token', token=token, _external=True)
            msg = Message('Password Reset Request',
                          sender=app.config['MAIL_USERNAME'],
                          recipients=[email])
            msg.body = f'To reset your password, visit the following link: {link}\n\nIf you did not make this request, please ignore this email.'
            try:
                mail.send(msg)
                flash('A password reset link has been sent to your email.', 'info')
            except Exception as e:
                app.logger.error(f"Password reset email failed: {e}")
                flash('Could not send reset link. Please try again later.', 'danger')
        else:
            flash('Email not found. Please check and try again.', 'danger')
        return redirect(url_for('login'))

    return render_template('reset_request.html')


@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_token(token):
    try:
        email = serializer.loads(token, salt='reset-salt', max_age=3600)
    except Exception:
        flash('The password reset link is invalid or has expired.', 'danger')
        return redirect(url_for('reset_request'))

    if request.method == 'POST':
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('reset_token', token=token))

        hashed_password = generate_password_hash(password)
        update_password(email, hashed_password)
        flash('Your password has been updated successfully. Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('reset_token.html', token=token)


@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if 'user_id' not in session:
        flash('Please log in to access the dashboard.', 'warning')
        return redirect(url_for('login'))

    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part in the request.', 'danger')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No file selected.', 'danger')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            unique_id = str(uuid.uuid4())
            upload_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{filename}")
            file.save(upload_path)

            conv_type = request.form.get('conversion_type')
            output_path = None
            try:
                if conv_type == 'pdf_to_word':
                    output_path = pdf_to_word(upload_path, unique_id)
                elif conv_type == 'pdf_to_jpg':
                    output_path = pdf_to_jpg_zip(upload_path, unique_id)
                elif conv_type == 'pdf_to_ipynb':
                    output_path = pdf_to_ipynb(upload_path, unique_id)
                else:
                    flash('Invalid conversion type selected.', 'danger')
                    return redirect(url_for('dashboard'))

                os.remove(upload_path)
                return send_converted_file(output_path)
            except Exception as e:
                app.logger.error(f"Conversion failed for {upload_path}: {e}")
                if upload_path and os.path.exists(upload_path): os.remove(upload_path)
                if output_path and os.path.exists(output_path): os.remove(output_path)
                flash('File conversion failed. Please try another file.', 'danger')
        else:
            flash('Invalid file type. Only PDF files are allowed.', 'danger')

    return render_template('dashboard.html', username=session.get('username'))
    
def send_converted_file(path):
    @after_this_request
    def cleanup(response):
        try:
            os.remove(path)
        except OSError as e:
            app.logger.error(f"Error removing file {path}: {e}")
        return response
    
    directory = os.path.dirname(path)
    filename = os.path.basename(path)
    return send_from_directory(directory, filename, as_attachment=True)


@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password']
        if username == ADMIN_USERNAME and check_password_hash(ADMIN_PASSWORD_HASH, password):
            session['admin'] = True
            flash('Admin login successful.', 'success')
            return redirect(url_for('admin_dashboard'))
        flash('Invalid admin credentials.', 'danger')
    return render_template('admin_login.html')

@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get('admin'):
        flash('You must be an admin to view this page.', 'danger')
        return redirect(url_for('admin_login'))

    with sqlite3.connect(DB_NAME) as conn:
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT id, username, email FROM users ORDER BY id DESC")
        users = c.fetchall()

    return render_template('admin_dashboard.html', users=users)

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin', None)
    flash('Admin logged out.', 'info')
    return redirect(url_for('admin_login'))


# --- PDF Conversion Functions ---
def pdf_to_word(pdf_path, unique_id):
    doc = Document()
    pdf_document = fitz.open(pdf_path)
    full_text = ""
    for page in pdf_document:
        full_text += page.get_text("text") + "\n"
    pdf_document.close()
    
    doc.add_heading('Converted from PDF', level=1)
    doc.add_paragraph(full_text)
    output_path = os.path.join(CONVERTED_FOLDER, f"{unique_id}_converted.docx")
    doc.save(output_path)
    return output_path

def pdf_to_jpg_zip(pdf_path, unique_id):
    pdf_document = fitz.open(pdf_path)
    image_folder = os.path.join(TEMP_FOLDER, unique_id)
    os.makedirs(image_folder, exist_ok=True)
    image_paths = []
    
    for i, page in enumerate(pdf_document):
        pix = page.get_pixmap(dpi=150)
        img_path = os.path.join(image_folder, f"page_{i+1}.jpg")
        pix.save(img_path)
        image_paths.append(img_path)
    pdf_document.close()
    
    zip_path = os.path.join(CONVERTED_FOLDER, f"{unique_id}_pages.zip")
    with zipfile.ZipFile(zip_path, 'w') as zf:
        for img in image_paths:
            zf.write(img, arcname=os.path.basename(img))
            os.remove(img)
            
    os.rmdir(image_folder)
    return zip_path

def pdf_to_ipynb(pdf_path, unique_id):
    pdf_document = fitz.open(pdf_path)
    full_text = "\n\n".join([page.get_text("text") for page in pdf_document])
    pdf_document.close()
    
    notebook = nbformat.v4.new_notebook()
    markdown_cell = nbformat.v4.new_markdown_cell(f"# Extracted Text from PDF\n\n---\n\n{full_text}")
    notebook['cells'] = [markdown_cell]
    
    output_path = os.path.join(CONVERTED_FOLDER, f"{unique_id}_converted.ipynb")
    with open(output_path, 'w', encoding='utf-8') as f:
        nbformat.write(notebook, f)
    return output_path


if __name__ == '__main__':
    init_db()
    app.run(debug=True)